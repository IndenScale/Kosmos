"""
通用工具函数模块，用于支持 Job 编排和执行。
"""
import os
import re
import uuid
import mimetypes
from io import BytesIO
from typing import List, Dict, Any

import docx
from docx.document import Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.text.paragraph import Paragraph
from docx.table import Table, _Cell

import olefile
import puremagic

from ...models.job import Job, JobType
# 延迟导入tasks模块以避免循环导入
# from ...tasks import (
#     process_document_actor,
#     chunk_document_actor,
#     analyze_figure_asset_job,
#     indexing_actor,
#     tagging_actor,
#     content_extraction_actor,
# )

def dispatch_job_actor(job: Job):
    """
    Dispatches a job to the appropriate Dramatiq actor based on its job type.
    This function now gets actors by name from the broker to avoid circular imports.
    """
    # Import the single source of truth for the broker
    from ...tasks.broker import broker
    
    actor_map = {
        JobType.DOCUMENT_PROCESSING: "process_document_actor",
        JobType.CONTENT_EXTRACTION: "content_extraction_actor",
        JobType.CHUNKING: "chunk_document_actor",
        JobType.ASSET_ANALYSIS: "analyze_asset_actor",
        JobType.INDEXING: "indexing_actor",
        JobType.TAGGING: "tagging_actor",
    }
    actor_name = actor_map.get(job.job_type)
    if actor_name:
        try:
            actor = broker.get_actor(actor_name)
            print(f"  - Dispatching job {job.id} (type: {job.job_type}) to actor '{actor_name}' on queue '{actor.queue_name}'")
            actor.send(str(job.id))
        except KeyError:
            print(f"Warning: Actor '{actor_name}' is not registered with the broker. Skipping job {job.id}.")
    else:
        print(f"Warning: No actor found for job type {job.job_type}")


def _get_embedded_object_type_name(mime_type: str) -> str:
    """
    根据MIME类型为内嵌对象生成一个描述性名称。
    """
    mime_type = mime_type.lower()
    if 'wordprocessingml' in mime_type or 'msword' in mime_type:
        return "Document"
    if 'spreadsheetml' in mime_type or 'ms-excel' in mime_type:
        return "Worksheet"
    if 'presentationml' in mime_type or 'ms-powerpoint' in mime_type:
        return "Presentation"
    if 'visio' in mime_type:
        return "Drawing"
    if 'chart' in mime_type:
        return "Chart"
    return "Object"

def replace_embedded_obj_with_text(doc: Document, rel_id: str, replace_text: str):
    """
    在 python-docx 的 Document 对象中，找到指定的嵌入对象并将其替换为一个文本段落。
    这是一个底层的 XML 操作。

    Args:
        doc: python-docx 的 Document 对象。
        rel_id: 嵌入对象的 a:blip relationship id。
        replace_text: 用于替换的文本。
    """
    doc_part = doc.part
    
    # 查找所有包含 <w:object> 的段落
    for p in doc.paragraphs:
        if f'r:id="{rel_id}"' in p._p.xml:
            # 这是一个复杂的情况，嵌入对象在段落内。
            # 最稳健的方法是清除该段落的所有内容，然后添加新文本。
            p.clear()
            p.add_run(replace_text)
            return

    # 查找所有包含 <w:object> 的表格单元格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # 递归地检查单元格内的段落
                for p in cell.paragraphs:
                    if f'r:id="{rel_id}"' in p._p.xml:
                        cell.text = replace_text
                        return

def decompose_docx_and_replace_ole(
    parent_filename: str,
    docx_content: bytes
) -> tuple[bytes, list[dict[str, Any]]]:
    """
    在内存中分解一个 .docx 文件，提取所有内嵌的文件，并用文本占位符替换它们。
    这是一个纯函数，没有数据库副作用。

    Args:
        parent_filename: 容器 DOCX 的文件名。
        docx_content: .docx 文件的字节内容。

    Returns:
        一个元组，包含：
        - 修改后的 docx 文件的字节内容。
        - 一个字典列表，每个字典代表一个内嵌文件。
    """
    embedded_files = []
    doc = docx.Document(BytesIO(docx_content))

    for rel_id, rel in doc.part.rels.items():
        if rel.reltype not in (
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/package",
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/oleObject"
        ):
            continue

        embedded_part = rel.target_part
        original_part_name = os.path.basename(embedded_part.partname)
        file_data = embedded_part.blob
        mime_type = embedded_part.content_type
        is_ole_unpacked = False

        if 'oleObject' in mime_type:
            try:
                with BytesIO(file_data) as ole_stream:
                    if olefile.isOleFile(ole_stream):
                        ole = olefile.OleFileIO(ole_stream)
                        stream_path = next((s for s in ['\x01Ole10Native', 'Package', 'CONTENTS'] if ole.exists(s)), None)
                        if stream_path:
                            stream = ole.openstream(stream_path)
                            if stream_path == '\x01Ole10Native': stream.seek(4)
                            extracted_data = stream.read()
                            if extracted_data:
                                file_data = extracted_data
                                is_ole_unpacked = True
            except Exception as e:
                print(f"警告: 解析OLE容器 '{original_part_name}' 失败: {e}")

        if is_ole_unpacked:
            try:
                # 尝试用 puremagic 从文件内容推断 MIME 类型
                magic_results = puremagic.magic_buffer(file_data)
                detected_mime = magic_results[0].mime_type if magic_results else "application/octet-stream"
                
                # 针对 PDF 的强化修复：如果 puremagic 失败，但文件头是 %PDF-，则强制设为 PDF
                if detected_mime == "application/octet-stream" and file_data.startswith(b'%PDF-'):
                    mime_type = "application/pdf"
                else:
                    mime_type = detected_mime

            except Exception:
                # 如果 puremagic 异常，也进行一次 PDF 文件头检查
                if file_data.startswith(b'%PDF-'):
                    mime_type = "application/pdf"
                else:
                    mime_type = "application/octet-stream"
        
        parent_name_base, _ = os.path.splitext(parent_filename)
        object_type_name = _get_embedded_object_type_name(mime_type)
        number_match = re.search(r'\d+', original_part_name)
        number_str = f"_{number_match.group(0)}" if number_match else ""
        extension = mimetypes.guess_extension(mime_type) or '.bin'
        final_filename = f"{parent_name_base}_Embedded_{object_type_name}{number_str}{extension}"

        # 生成唯一的占位符ID，用于后续替换
        placeholder_id = str(uuid.uuid4())

        embedded_files.append({
            "placeholder_id": placeholder_id,
            "rel_id": rel_id,
            "filename": final_filename,
            "content": file_data,
            "mime_type": mime_type,
            "size": len(file_data)
        })

    # 在所有信息收集完毕后，再进行替换操作
    for file_info in embedded_files:
        # 这里的文本是临时的，最终会被替换为包含 document_id 的标签
        replace_embedded_obj_with_text(doc, file_info['rel_id'], f"[[KOSMOS_EMBED_PLACEHOLDER:{file_info['placeholder_id']}]]")

    # 将修改后的文档保存到内存中
    modified_stream = BytesIO()
    doc.save(modified_stream)
    modified_docx_content = modified_stream.getvalue()

    return modified_docx_content, embedded_files
